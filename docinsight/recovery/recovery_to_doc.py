import os

from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from table_process import HtmlToDocx

import logging

logger = logging.getLogger("docinsight.recovery")

def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = shared.Pt(6.5)

    flag = 1  # 当前布局标记
    previous_layout = None  # 用于记录前一个文本框的布局

    for i, region in enumerate(res):
        if len(region["res"]) == 0:
            continue

        # 检查当前布局是否发生变化，避免重复创建相同的布局
        current_layout = region["layout"]
        if current_layout != previous_layout:
            if current_layout == "single":
                section = doc.add_section(WD_SECTION.CONTINUOUS)
                section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")
                flag = 1
            elif current_layout == "double":
                section = doc.add_section(WD_SECTION.CONTINUOUS)
                section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "2")
                flag = 2
            elif current_layout == "triple":
                section = doc.add_section(WD_SECTION.CONTINUOUS)
                section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "3")
                flag = 3
            previous_layout = current_layout  # 更新前一个布局记录

        # 根据区域类型插入内容
        if region["type"].lower() == "figure":
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(
                excel_save_folder, "{}_{}.jpg".format(region["bbox"], region["img_idx"])
            )
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            # 插入图片，宽度根据列布局决定
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2.5))
            elif flag == 3:
                run.add_picture(img_path, width=shared.Inches(1.5))

        elif region["type"].lower() == "title":
            doc.add_heading(region["res"][0]["text"])

        elif region["type"].lower() == "table":
            parser = HtmlToDocx()
            parser.table_style = "TableGrid"
            parser.handle_table(region["res"]["html"], doc)

        else:  # 默认处理文本区域
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for j, line in enumerate(region["res"]):
                if j == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)  # 首行缩进
                text_run = paragraph.add_run(line["text"] + " ")
                text_run.font.size = shared.Pt(10)

    # 保存为 docx 文件
    docx_path = os.path.join(save_folder, "{}_ocr.docx".format(img_name))
    doc.save(docx_path)
    logger.info("docx saved to {}".format(docx_path))


def sorted_layout_boxes(res, w):
    """
    根据文本框的分布自适应排序，支持一竖版、两竖版和三竖版的布局，考虑跨度大的列
    args:
        res(list): ppstructure结果
        w(int): 文档宽度
    return:
        排序后的结果(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]["layout"] = "single"
        return res

    # 先根据y坐标从上到下，再根据x坐标从右到左排序
    sorted_boxes = sorted(res, key=lambda x: (x["bbox"][1], -x["bbox"][0]))
    _boxes = list(sorted_boxes)

    res_left = []
    res_center = []
    res_right = []
    new_res = []

    column_thresholds = [w / 3, 2 * w / 3]
    tolerance = 0.02 * w  # 设置一个容差，用于判断是否接近列边界

    # 第一轮：分类列，确定各个列的框分布
    for i in range(num_boxes):
        current_box = _boxes[i]
        box_left, box_right = current_box["bbox"][0], current_box["bbox"][2]
        box_width = box_right - box_left

        # 判断列布局，确保每个box只分配到一个列
        if box_width > column_thresholds[1]:
            current_box["layout"] = "spanning"
            new_res.append(current_box)
        elif box_right < column_thresholds[0] + tolerance:  # 左列
            res_left.append(current_box)
        elif box_left > column_thresholds[1] - tolerance:  # 右列
            res_right.append(current_box)
        elif column_thresholds[0] - tolerance <= box_left <= column_thresholds[1] + tolerance:  # 中列
            res_center.append(current_box)
        else:
            res_left.append(current_box)

    # 第二轮：根据列分布情况，判断具体的布局
    for box in res_left:
        if res_center and res_right:
            box["layout"] = "triple"
        elif res_right:
            box["layout"] = "double"
        elif res_center:
            box["layout"] = "double"
        else:
            box["layout"] = "single"
        new_res.append(box)

    for box in res_center:
        box["layout"] = "triple" if res_left and res_right else "double"
        new_res.append(box)

    for box in res_right:
        box["layout"] = "triple" if res_center else "double"
        new_res.append(box)

    _ = new_res

    return new_res
